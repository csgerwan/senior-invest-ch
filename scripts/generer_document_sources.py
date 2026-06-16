"""
Génère un document Word (.docx) « Sources & Méthodologie » avec liens cliquables.
Sortie : ~/Desktop/Senior_Invest_CH_Sources_et_Methodologie.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

OUT = Path.home() / "Desktop" / "Senior_Invest_CH_Sources_et_Methodologie.docx"
BLEU = RGBColor(0x1F, 0x4E, 0x79)


def hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rpr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run)
    paragraph._p.append(link)


def para(doc, text="", bold=False, size=11, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p


def source(doc, titre, desc, liens):
    """Un bloc source : titre en gras, description, puis liens cliquables."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("• " + titre); r.bold = True; r.font.size = Pt(11)
    if desc:
        pd = doc.add_paragraph(); pd.paragraph_format.space_after = Pt(2)
        pd.paragraph_format.left_indent = Pt(14)
        rd = pd.add_run(desc); rd.italic = True; rd.font.size = Pt(10)
    for libelle, url in liens:
        pl = doc.add_paragraph(); pl.paragraph_format.space_after = Pt(2)
        pl.paragraph_format.left_indent = Pt(14)
        pl.add_run(libelle + " : ").font.size = Pt(10)
        hyperlink(pl, url, url)


def h1(doc, text):
    h = doc.add_heading(level=1)
    r = h.add_run(text); r.font.color.rgb = BLEU


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Titre
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rt = t.add_run("Senior Invest CH"); rt.bold = True; rt.font.size = Pt(22)
    rt.font.color.rgb = BLEU
    st = doc.add_paragraph()
    rs = st.add_run("Sources des données & méthodologie des calculs")
    rs.font.size = Pt(13); rs.italic = True
    sub = doc.add_paragraph()
    sub.add_run("Outil d'aide à la décision — investissement en hébergement senior, "
                "canton de Vaud").font.size = Pt(10)
    dt = doc.add_paragraph()
    dt.add_run("Document de référence — juin 2026").font.size = Pt(9)
    pcode = doc.add_paragraph(); pcode.add_run("Code source du projet : ").font.size = Pt(9)
    hyperlink(pcode, "https://github.com/csgerwan/senior-invest-ch",
              "github.com/csgerwan/senior-invest-ch")

    para(doc, "")
    intro = doc.add_paragraph()
    intro.add_run(
        "Toutes les données utilisées sont publiques, officielles, sourcées et datées. "
        "Chaque commune vaudoise est identifiée par son numéro OFS, clé qui relie "
        "l'ensemble des jeux de données entre eux. Ce document liste les sources exactes "
        "(liens cliquables) puis détaille les calculs effectués.").font.size = Pt(11)

    # ----- SOURCES -----
    h1(doc, "1. Sources des données")

    para(doc, "Démographie & géographie", bold=True, size=12, space_after=2)
    source(doc, "Population résidante permanente par âge et par commune (2024)",
           "Office fédéral de la statistique (OFS) — STAT-TAB, table px-x-0102010000_103.",
           [("Table OFS", "https://www.pxweb.bfs.admin.ch/pxweb/fr/px-x-0102010000_103/"),
            ("Portail population OFS",
             "https://www.bfs.admin.ch/bfs/fr/home/statistiques/population.html")])
    source(doc, "Limites des communes vaudoises",
           "Découpage officiel OFS / swisstopo, diffusé via Opendatasoft.",
           [("Jeu de données communes",
             "https://public.opendatasoft.com/explore/dataset/georef-switzerland-gemeinde/"),
            ("Géoportail fédéral", "https://www.geo.admin.ch")])

    para(doc, "Concurrence — EMS (établissements médico-sociaux)", bold=True, size=12, space_after=2)
    source(doc, "Liste officielle des EMS vaudois et nombre de lits (2024)",
           "Canton de Vaud — Liste LAMal des EMS et divisions C (mandat 2024).",
           [("PDF officiel vd.ch",
             "https://www.vd.ch/fileadmin/user_upload/themes/social/EMS/Professionnel/Tarifs_soho/Liste_LAMal_2024.pdf")])
    source(doc, "Nombre de lits cantonal & taux d'occupation",
           "INFOSAN — chiffres-clés de la santé vaudoise (6 986 lits ; occupation ~98 %).",
           [("INFOSAN — EMS, lits", "https://infosan.vd.ch/chiffres-cles/ems-lits")])
    source(doc, "Géolocalisation des EMS",
           "Coordonnées et adresses recoupées via trois sources publiques.",
           [("OpenStreetMap", "https://www.openstreetmap.org"),
            ("API officielle geo.admin", "https://api3.geo.admin.ch"),
            ("Annuaire search.ch", "https://tel.search.ch")])

    para(doc, "Pouvoir d'achat", bold=True, size=12, space_after=2)
    source(doc, "Revenu par commune (impôt fédéral direct, 2022)",
           "Administration fédérale des contributions (AFC/ESTV) — statistique IFD des "
           "personnes physiques, par commune.",
           [("ESTV — IFD par commune",
             "https://www.estv.admin.ch/estv/fr/accueil/afc/statistiques-fiscales/statistiques-fiscales-general/statistiques-impot-federal-direct/ifd-pp-communes-depuis-1983.html")])
    source(doc, "Charge fiscale — coefficient d'impôt communal (2026)",
           "Canton de Vaud — arrêtés d'imposition / tableaux des impôts communaux.",
           [("vd.ch — taux des impôts communaux",
             "https://www.vd.ch/etat-droit-finances/communes/finances-communales/arretes-dimposition-et-tableaux-des-impots-communaux")])
    source(doc, "Fortune des personnes physiques (par district, 2022)",
           "Statistique Vaud — recettes fiscales cantonales (table T18.02.16).",
           [("StatVD — recettes fiscales",
             "https://www.vd.ch/etat-droit-finances/statistique/statistiques-par-domaine/18-finances-publiques-et-administration/recettes-fiscales-cantonales")])

    para(doc, "Immobilier", bold=True, size=12, space_after=2)
    source(doc, "Prix au m² (annonces, 2026)",
           "Annonces d'achat de Homegate.ch (premier portail immobilier suisse), "
           "extraites via la plateforme Apify. Prix demandés, et non prix de transaction.",
           [("Homegate.ch", "https://www.homegate.ch"),
            ("Actor Apify (Homegate)",
             "https://apify.com/santamaria-automations/homegate-scraper")])

    # ----- CALCULS -----
    doc.add_page_break()
    h1(doc, "2. Méthodologie des calculs")

    para(doc, "2.1 Démographie", bold=True, size=12, space_after=2)
    para(doc, "• Population 65+ = somme des classes d'âge 65-69, 70-74, …, 100 ans et plus. "
              "Idem pour les 80+ (classes 80-84 … 100+).", size=10)
    para(doc, "• Part des seniors (%) = population de la tranche ÷ population totale × 100.", size=10)

    para(doc, "2.2 Concurrence EMS", bold=True, size=12, space_after=2)
    para(doc, "• Lits par commune = somme des lits des EMS situés dans la commune "
              "(chaque EMS rattaché à sa commune par ses coordonnées GPS).", size=10)
    para(doc, "• « Éloignement des EMS » = distance, en km, entre le centre de la commune "
              "et l'EMS le plus proche (calcul géographique sur coordonnées projetées). "
              "Une distance élevée signale une zone mal desservie.", size=10)
    para(doc, "• Taux d'occupation : valeur cantonale officielle (~98 %), non disponible "
              "par établissement.", size=10)

    para(doc, "2.3 Pouvoir d'achat", bold=True, size=12, space_after=2)
    para(doc, "• Revenu net médian estimé : l'OFS/ESTV fournit le nombre de contribuables "
              "par tranche de revenu. On en déduit la médiane par la formule des classes :", size=10)
    pf = doc.add_paragraph(); pf.paragraph_format.left_indent = Pt(28)
    rf = pf.add_run("médiane = L + ((N/2 − F) / f) × largeur_de_classe"); rf.italic = True; rf.font.size = Pt(10)
    para(doc, "où L = borne basse de la classe contenant le médian, N = total des "
              "contribuables, F = effectif cumulé avant cette classe, f = effectif de la classe.", size=10)
    para(doc, "• Part de hauts revenus = % de contribuables déclarant plus de 100 000 CHF.", size=10)
    para(doc, "• Fortune par district = (part du district dans la fortune cantonale ÷ part "
              "du district dans les contribuables) × fortune moyenne cantonale par contribuable.", size=10)
    para(doc, "• Indice de pouvoir d'achat (0-100) = moyenne de deux composantes normalisées : "
              "le revenu (plus il est haut, mieux c'est) et la charge fiscale inversée "
              "(plus le coefficient d'impôt est bas, mieux c'est).", size=10)

    para(doc, "2.4 Immobilier", bold=True, size=12, space_after=2)
    para(doc, "• Prix au m² d'une annonce = prix de vente ÷ surface habitable.", size=10)
    para(doc, "• Surface : lue dans le texte de l'annonce ; si absente, estimée via le nombre "
              "de pièces (ratio médian observé ≈ 29,6 m² par pièce).", size=10)
    para(doc, "• Prix au m² d'une commune = moyenne (et médiane) des annonces de la commune, "
              "après filtrage des valeurs aberrantes (fourchette 3 000 – 30 000 CHF/m²).", size=10)
    para(doc, "• Communes sans bien en vente : prix estimé = médiane des 3 communes "
              "géographiquement les plus proches disposant d'un prix.", size=10)

    para(doc, "2.5 Score d'opportunité (0-100)", bold=True, size=12, space_after=2)
    para(doc, "Croisement de 4 critères, chacun ramené à une note de 0 à 100, puis "
              "moyenne pondérée (poids ajustables dans l'outil) :", size=10)
    for txt in [
        "1. Part de seniors (80+) élevée  →  forte demande potentielle",
        "2. Éloignement des EMS existants  →  zone mal desservie",
        "3. Pouvoir d'achat élevé  →  capacité des résidents à financer",
        "4. Prix immobilier au m² bas  →  foncier abordable (note inversée)",
    ]:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(16)
        p.add_run(txt).font.size = Pt(10)
    para(doc, "Normalisation « min-max » utilisée partout : note = (valeur − minimum) ÷ "
              "(maximum − minimum) × 100.", size=10)

    # ----- LIMITES -----
    h1(doc, "3. Limites & précautions")
    for txt in [
        "Prix immobiliers = prix demandés (annonces), pas des prix de transaction réels ; "
        "à valider par une source professionnelle (Wüest Partner, FPRE) avant engagement.",
        "Fortune disponible au niveau district seulement (secret fiscal communal) : toutes "
        "les communes d'un même district partagent la valeur.",
        "Revenu (impôt fédéral direct) : sous-estime légèrement les bas revenus et porte sur "
        "l'année 2022.",
        "Les estimations (surface par pièces, prix par voisinage) sont signalées par une "
        "colonne « fiabilité » dans l'outil.",
        "Outil d'aide à la décision : les pondérations du score sont des hypothèses à ajuster "
        "selon la stratégie d'investissement.",
    ]:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.add_run("• " + txt).font.size = Pt(10)

    doc.save(OUT)
    print(f"✅ Document généré : {OUT}")


if __name__ == "__main__":
    main()
